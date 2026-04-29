from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "AeroSlot_Dashboard_System_Complete_Explanation.docx"
ARCH = ROOT / "System_Architecture_Diagram.jpg"
CLASS = ROOT / "C++_Class_Diagram.jpg"


def set_cell_text(cell, text):
    cell.text = str(text)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h)
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "20")
        elem.set(qn("w:color"), "8EC5DA")
        pgBorders.append(elem)
    sectPr.append(pgBorders)


def setup(doc):
    for sec in doc.sections:
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)
        border(sec)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Calibri"
        styles[name].font.bold = True


def title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AeroSlot Airport Runway Scheduler")
    r.bold = True
    r.font.size = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Complete Dashboard, System Flow, Algorithms, ML, C++, Database, Dataset, and Viva Explanation")
    r.bold = True
    r.font.size = Pt(13)
    doc.add_paragraph("This document explains the current simplified dashboard and the project logic in simple teacher-friendly language.")
    doc.add_paragraph("Current dashboard tabs: Flight Intake, Simulation Control, and Runway Status.")
    doc.add_paragraph("Project database explanation: MySQL is used as the final project database for storing flights, predictions, schedules, runway states, events, emergencies, and algorithm switch logs.")
    if ARCH.exists():
        doc.add_picture(str(ARCH), width=Inches(6.5))
    doc.add_page_break()


def dashboard_section(doc):
    doc.add_heading("1. Dashboard Overview", level=1)
    doc.add_paragraph("The dashboard is the main user interface of the Airport Runway Scheduler. It allows the user to add flights, load a dataset, select/run scheduling algorithms, and view live runway status.")
    add_table(doc, ["Dashboard Part", "What It Means", "Main File"], [
        ["Top Navigation Bar", "Shows project name and three tabs: Flight Intake, Simulation Control, Runway Status.", "frontend/src/components/FlightInputEngine.jsx"],
        ["Flight Intake Tab", "Used to enter required flight details or load dataset/demo flights.", "FlightInputEngine.jsx and useTelemetryStore.js"],
        ["Simulation Control Tab", "Used to select algorithm, run scheduler, optimize runways, inject emergency, or inject storm.", "FlightInputEngine.jsx and useTelemetryStore.js"],
        ["Runway Status Tab", "Shows total/free/occupied runways and live runway state board.", "FlightInputEngine.jsx and useTelemetryStore.js"],
        ["Background and Animations", "Aviation background, aircraft motion, algorithm motion, runway state animations.", "frontend/src/index.css"],
    ])

    doc.add_heading("2. Flight Intake Tab: Every Input and Card", level=1)
    doc.add_paragraph("This tab takes only the required details needed for runway scheduling. The system avoids unnecessary inputs so that a normal user can submit flights easily.")
    add_table(doc, ["Input / Component", "Meaning", "Options / Example", "How It Affects Scheduling"], [
        ["Flight Number", "Unique code of the flight.", "AA772, AI101, EM901", "Used as the identity in queues, logs, schedule results, and runway cards."],
        ["Airline", "Airline or operator name.", "Air India, IndiGo, Vistara, SpiceJet, Akasa Air, American, United, Delta, Medevac", "Used for display and database history. It does not directly decide priority unless the flight is emergency/demo."],
        ["Operation", "Whether flight is landing or takeoff.", "Landing or Takeoff", "Landing flights go to Landing Queue. Takeoff flights go to Takeoff Queue. It also affects model features such as altitude/velocity in Python API."],
        ["Wake Size", "Aircraft wake turbulence category.", "Light, Medium, Heavy", "Used for wake separation. Heavy aircraft normally need more separation than Medium or Light."],
        ["From", "Origin airport.", "DEL, BOM, BLR, HYD, etc.", "Used in route display, history, and runway switch message."],
        ["To", "Destination airport.", "DEL, BOM, BLR, HYD, etc.", "Used in route display and history."],
        ["Schedule Time (sec)", "Expected arrival/takeoff time in simulation seconds.", "120, 180, 300, 450, 600, 900, 1200, 1800 sec", "This is the base time used for slot assignment. Lower time means the flight is expected earlier."],
        ["Runway Use Time (sec)", "How long the aircraft occupies the runway.", "35, 45, 50, 60, 75, 90 sec", "Used as runway timer. Longer use time keeps runway occupied longer and delays next aircraft."],
        ["Flight Priority", "Passenger/flight importance.", "Normal, VIP, Medical, High Density", "Adds weight to priority score. Medical/VIP/High Density flights receive more importance than Normal."],
        ["Cargo Need", "Cargo urgency type.", "Routine, Perishable, Medical, Military", "Adds cargo weight to priority score. Medical/Military cargo gets higher weight."],
        ["Fuel Urgency", "Fuel urgency from 1 to 10.", "Slider 1-10", "Higher fuel urgency increases priority score and reduces estimated delay in fallback formula."],
        ["Safety Gap", "Minimum safe separation in seconds.", "Example: 90 sec", "Used as separation buffer before scheduling next aircraft on a runway."],
        ["Preferred Runway", "Manual runway request.", "AUTO, 27L, 27R, 09L, 09R", "If Auto Runway Assignment is enabled, scheduler chooses runway. If disabled, requested runway is used when valid."],
        ["Runway Condition", "Current runway surface condition.", "DRY, WET, CONTAMINATED", "Used as operational context. Wet/contaminated runway is explained as riskier runway condition."],
        ["Medical Distress", "Emergency medical situation.", "Checkbox", "Triggers emergency logic. Emergency flights are considered first and can switch algorithm to Preemptive Emergency."],
        ["Technical Distress", "Technical problem in aircraft.", "Checkbox", "Also treated as emergency. It increases priority and may trigger emergency scheduling."],
        ["Auto Runway Assignment", "Allows system to choose runway automatically.", "Checkbox", "When checked, scheduler chooses the runway with lower load. This is the recommended option."],
        ["Visibility (km)", "How far pilot/tower can see.", "3, 5, 8, 10, 12, 15 km", "Low visibility increases weather risk and delay. Very low visibility can push system toward Hybrid Adaptive AI."],
        ["Storm Severity", "Weather storm level.", "0 to 10", "High storm increases delay/risk. If storm is high, the algorithm can switch to Hybrid Adaptive AI."],
        ["Traffic Level", "Current airport congestion.", "LOW, MEDIUM, HIGH", "HIGH traffic can switch scheduler to Greedy Min Delay because runway load is heavy."],
    ])

    add_table(doc, ["Button / Card", "What It Does", "Where Logic Is Stored"], [
        ["Add Flight", "Adds one manually entered flight to active flights, predicts basic delay/risk, and logs FLIGHT_ADDED.", "addFlight() in frontend/src/store/useTelemetryStore.js"],
        ["Batch Load Dataset", "Loads multiple sample/dataset flights into the existing list without deleting already added flights.", "batchLoadDataset() and loadDemoTraffic() in useTelemetryStore.js"],
        ["Load Demo Dataset", "Adds demo flights including normal, takeoff, heavy, and emergency examples.", "loadDemoTraffic() in useTelemetryStore.js"],
        ["Estimated Delay Card", "Shows expected delay for latest added flight.", "estimatePrediction() / predictionResult in useTelemetryStore.js"],
        ["Conflict Risk Card", "Shows calculated safety/conflict risk percentage.", "estimatePrediction() and local_prediction()"],
        ["Priority Score Card", "Shows numeric priority score calculated from emergency/fuel/passenger/cargo fields.", "priorityScore() in useTelemetryStore.js"],
        ["Recommended Runway Card", "Shows suggested runway from automatic logic or backend prediction.", "estimatePrediction() / recommendedRunway"],
        ["Landing Queue", "Shows all current landing flights.", "QueueMonitor component"],
        ["Takeoff Queue", "Shows all current takeoff flights.", "QueueMonitor component"],
        ["Recently Added Flights", "Shows recent flights with flight ID, route, and status.", "FlightHistoryMini component"],
    ])

    doc.add_heading("2.1 What Happens When Inputs Change?", level=2)
    add_table(doc, ["Input Change", "System Reaction"], [
        ["Storm Severity increases", "Weather penalty increases. Delay/risk increases. If storm >= 7 or risk is high, Hybrid Adaptive AI may be selected."],
        ["Traffic Level becomes HIGH", "Congestion penalty increases. Scheduler may select Greedy Min Delay to reduce total delay."],
        ["Medical/Technical Distress checked", "Emergency flag becomes true. Scheduler selects Preemptive Emergency and places emergency aircraft first."],
        ["Fuel Urgency increases", "Priority score increases and delay formula gives urgency credit."],
        ["Wake Size becomes Heavy", "Wake penalty/separation increases, making scheduling more careful."],
        ["Runway Use Time increases", "Runway remains occupied longer, so queue clears slower."],
        ["Schedule Time earlier", "Flight gets earlier ETA and may be scheduled before later aircraft, depending on algorithm."],
        ["Auto Runway Assignment checked", "System balances runway load automatically."],
        ["Preferred Runway selected manually", "Flight tries to use that runway if auto assignment is disabled."],
    ])


def simulation_section(doc):
    doc.add_heading("3. Simulation Control Tab", level=1)
    doc.add_paragraph("This tab controls the scheduling run. It does not remove algorithm logic; it simply lets the user select or trigger scheduling actions.")
    add_table(doc, ["Component", "Meaning", "What Happens"], [
        ["Algorithm Cards", "Shows FCFS, Priority Scheduling, Preemptive Emergency, Greedy Min Delay, and Hybrid Adaptive AI.", "Clicking a card selects preferred algorithm for the next run. Actual active algorithm is shown in Algorithm Flow."],
        ["Run Scheduler", "Runs the scheduler on current flights.", "System checks emergency, congestion, risk, storm, and selected algorithm. Then it assigns runways and slot times."],
        ["Optimize Runways", "Forces Greedy Min Delay.", "Schedules flights using greedy delay minimization and updates runway board."],
        ["Inject Emergency", "Adds emergency flight.", "Creates emergency flight, adds it to queue, and switches active algorithm to Preemptive Emergency."],
        ["Inject Storm", "Creates bad weather condition.", "Sets storm to 9, traffic to HIGH, visibility to 3 km, and runway condition to WET."],
        ["Status Banner", "Black message bar.", "Shows latest scheduler status such as selected algorithm or emergency/storm event."],
        ["Algorithm Flow", "Visualization of active algorithm.", "Highlights the real active algorithm. Airplane icon moves to that algorithm block."],
        ["Scheduled Flight Logs", "Table and event list.", "Shows flight ID, runway, slot time, delay, and recent scheduling events."],
    ])

    doc.add_heading("4. Runway Status Tab", level=1)
    add_table(doc, ["Card / Component", "Meaning", "How It Updates"], [
        ["Total Runways", "Total available runways in system.", "Current project has 4 runways: 27L, 27R, 09L, 09R."],
        ["Free Runways", "Number of runways with no active aircraft.", "Updates every second based on runway timers and queues."],
        ["Occupied Runways", "Number of runways currently in use.", "Increases after scheduling; decreases when runway timer reaches 0 and queue is empty."],
        ["Runway Animation Strip", "Visual aircraft movement from approach to departure.", "Runs continuously to show runway operation."],
        ["Runway State Board", "Main live runway cards.", "Each card shows runway name, status, active flight, queue count, timer, wake gap, and utilization."],
        ["Runway Status", "FREE or OCCUPIED.", "Changes to FREE when timer reaches 0 and no queued aircraft remains."],
        ["Active Flight", "Current flight occupying the runway.", "Changes to next queued flight when timer finishes."],
        ["Queue", "Number of flights assigned to that runway.", "Decreases when one flight completes."],
        ["Timer", "Remaining runway occupancy time.", "Counts down every second."],
        ["Wake", "Safety separation requirement.", "Comes from Safety Gap or wake category."],
        ["Util", "Simple utilization percentage.", "Higher when more flights are assigned to that runway."],
        ["Runway Notice", "Temporary blue message.", "Displays for 5 seconds when next flight is scheduled to a runway."],
    ])


def flow_section(doc):
    doc.add_heading("5. Full System Flow From Start to Storage", level=1)
    numbered(doc, [
        "User opens the dashboard. React loads App.jsx, which renders FlightInputEngine.jsx.",
        "Zustand store useTelemetryStore.js initializes default state: algorithm AUTO, selectedAlgorithm AUTO, empty flights, empty runway states, default weather LOW congestion.",
        "Dashboard initially shows Flight Intake tab. Active scheduling algorithm visualization treats AUTO as FCFS because FCFS is the normal default baseline.",
        "User enters flight data or clicks Batch Load Dataset / Load Demo Dataset.",
        "When Add Flight is clicked, addFlight() creates a complete flight object, calculates prediction using estimatePrediction(), saves it in flights array, and sends data to /api/flights if backend is running.",
        "Landing flights appear in Landing Queue. Takeoff flights appear in Takeoff Queue. Recently Added Flights table also updates.",
        "User opens Simulation Control and clicks Run Scheduler.",
        "runSchedulerAPI() calls chooseAlgorithm() to decide the actual algorithm based on emergency, congestion, risk, storm, and selected algorithm.",
        "scheduleFlights() orders flights according to the selected algorithm and assigns each flight to a runway.",
        "runwaysFromSchedule() groups scheduled flights by runway and creates live runway states with queue, active aircraft, timer, wake gap, and utilization.",
        "The schedule is posted to /api/schedule if the Flask backend is running.",
        "Python api_server.py receives schedule and db_service.py stores schedule_results, delay_predictions, runway_states, event_history, and algorithm_switch_log in database tables.",
        "Runway Status tab updates every second using tickRunways(). When a runway timer reaches zero, the next queued flight becomes active or the runway becomes free.",
        "A temporary 5-second message shows when a next flight is scheduled to a runway."
    ])
    add_table(doc, ["Process", "Frontend File", "Backend / C++ / DB File"], [
        ["Dashboard rendering", "frontend/src/main.jsx, App.jsx, FlightInputEngine.jsx", "-"],
        ["State and scheduling logic", "frontend/src/store/useTelemetryStore.js", "-"],
        ["Styling and animation", "frontend/src/index.css", "-"],
        ["Prediction API", "useTelemetryStore.js fetch /api/predict", "python_ai/api_server.py"],
        ["Flight save", "safePost('/flights')", "api_server.py -> db_service.py -> database"],
        ["Schedule save", "safePost('/schedule')", "api_server.py -> db_service.py -> schedule_results"],
        ["Algorithm switch log", "safePost('/switches')", "api_server.py -> db_service.py -> algorithm_switch_log"],
        ["C++ scheduling concepts", "Shown in UI logic and docs", "cpp_scheduler and SchedulerEngine headers"],
    ])


def algorithm_section(doc):
    doc.add_heading("6. Scheduling Algorithms: Logic, Inputs, and Formulas", level=1)
    doc.add_paragraph("Default baseline algorithm: FCFS. In the frontend state, algorithm initially starts as AUTO, but the dashboard treats AUTO as FCFS until conditions cause a switch.")
    add_table(doc, ["Algorithm", "When It Is Used", "Main Inputs Considered", "Main Logic / Formula"], [
        ["FCFS", "Normal stable traffic with no emergency/high congestion/high risk.", "Schedule Time, input order, runway availability.", "Flights are handled in request order. Baseline idea: first flight added/expected first is scheduled first."],
        ["Priority Scheduling", "When user selects priority mode or medium load style priority is required.", "Emergency, fuel urgency, passenger priority, cargo need, waiting time.", "priorityScore = emergency + fuelUrgency*9 + passengerWeight + cargoWeight. Higher score comes first."],
        ["Preemptive Emergency", "When any flight has medical/technical/fuel emergency.", "Emergency flag, priority score, schedule time.", "Emergency flights are sorted first: emergencyWeight desc, then priorityScore desc."],
        ["Greedy Min Delay", "When congestion is HIGH, flights >= 8, or Optimize Runways is clicked.", "Schedule time, runway use time, runway load.", "Sorts by ETA + runwayOccupancy. Also C++ concept uses penalty = separation*2 - priority."],
        ["Hybrid Adaptive AI", "When risk is high or storm severity is high.", "Priority score, ETA, ML/fallback risk, storm, congestion.", "Sorts by priorityScore / max(1, ETA). High priority and earlier time get advantage."],
    ])
    add_table(doc, ["Formula", "Meaning"], [
        ["priorityScore = emergency + fuelUrgency*9 + passengerWeight + cargoWeight", "Calculates flight importance."],
        ["delay = max(5, 45 + congestion + weather + wake - urgencyCredit)", "Frontend fallback delay estimation."],
        ["risk = clamp(0.08, 0.96, delay/150 + congestionRisk + emergencyRisk)", "Conflict/risk percentage estimation."],
        ["slotTime = max(ETA, runwayLoad[runway] + separation)", "Finds safe time slot for runway."],
        ["scheduledDelay = max(0, slotTime - ETA) + risk*8", "Final displayed delay after scheduling."],
        ["runwayLoad = slotTime + runwayOccupancy", "Updates when runway becomes available again."],
        ["wakeBuffer = baseWakeMatrix[leader, follower] * weatherFactor", "C++ wake separation safety concept."],
        ["interval overlap if max(start1,start2) < min(end1,end2)", "C++ conflict detector checks if two runway slots overlap."],
    ])
    doc.add_paragraph("Algorithm switching rule in simple language:")
    numbered(doc, [
        "If emergency exists, use Preemptive Emergency.",
        "Else if traffic is HIGH or flight count is 8 or more, use Greedy Min Delay.",
        "Else if conflict risk is high or storm is high, use Hybrid Adaptive AI.",
        "Else use selected algorithm, and if selection is AUTO, use FCFS."
    ])


def concepts_section(doc):
    doc.add_heading("7. Object Oriented Programming and OS Concepts", level=1)
    if CLASS.exists():
        doc.add_picture(str(CLASS), width=Inches(6.5))
    add_table(doc, ["Concept", "Where Used", "Simple Explanation"], [
        ["Class / Struct", "Aircraft, Runway, Scheduler, ConflictDetector, Strategy classes", "Represents real-world objects like aircraft and runway."],
        ["Encapsulation", "Runway and RunwayResourceManager", "Runway state is changed using methods instead of direct unsafe changes."],
        ["Inheritance", "ISchedulerStrategy -> FCFSStrategy, PriorityStrategy, GreedyStrategy", "Different algorithms follow the same interface."],
        ["Polymorphism", "Scheduler uses unique_ptr<ISchedulerStrategy>", "Scheduler can call the same method while actual algorithm changes."],
        ["Abstraction", "ISchedulerStrategy, IEmergencyObserver", "Only important behavior is exposed, internal logic is hidden."],
        ["Operator Overloading", "Aircraft::operator<", "Defines how aircraft are compared in priority structures."],
        ["STL Queue", "FCFSStrategy", "Maintains first-come-first-served order."],
        ["Priority Queue / Heap", "PriorityStrategy, GreedyStrategy", "Gets highest priority or lowest delay candidate quickly."],
        ["Map", "WakeSeparationManager, ConflictDetector", "Stores wake matrices and runway interval schedules."],
        ["Smart Pointer unique_ptr", "Scheduler, ConstraintManager", "Automatic memory management for strategy objects."],
        ["Mutex / lock_guard", "RunwayResourceManager", "Operating system concurrency concept; prevents simultaneous unsafe runway updates."],
        ["Optional", "RunwayResourceManager", "Represents whether a runway has an aircraft or not."],
        ["Timer / Scheduling", "tickRunways() in frontend store", "Simulates OS-like periodic updates every second."],
        ["Process orchestration", "launch_system.py", "Starts Python service, C++ engine, and React frontend."],
        ["Inter-process communication idea", "Frontend fetch API + Flask API", "Frontend and backend communicate using HTTP requests."],
    ])


def ml_tech_section(doc):
    doc.add_heading("8. ML Model Used and How It Works", level=1)
    add_table(doc, ["Model", "File", "Input Parameters", "Output", "Purpose"], [
        ["Random Forest Regressor", "python_ai/pipeline/train_rf.py, python_ai/models/rf_delay_model.pkl", "geoaltitude, velocity, wake_factor, airspace_congestion, occupancy_est_sec", "predicted delay in seconds", "Predicts expected runway/flight delay."],
        ["XGBoost Priority Classifier", "python_ai/pipeline/train_xgboost.py, xgb_priority_model.pkl", "Engineered flight and operation features", "priority class", "Supports priority decision intelligence."],
        ["XGBoost Congestion Regressor", "python_ai/pipeline/train_xgboost.py, xgb_congestion_model.pkl", "Congestion and feature vectors", "congestion estimate", "Supports congestion-aware scheduling."],
        ["Fallback Prediction", "estimatePrediction() in useTelemetryStore.js and local_prediction() in api_server.py", "wake, congestion, storm, visibility, fuel urgency, emergency", "delay, risk, score, confidence, recommended runway", "Keeps demo working even when backend/model is unavailable."],
    ])
    doc.add_paragraph("Random Forest works by creating many decision trees and averaging their results. It is good for tabular project data because each tree learns different rules from the features.")
    doc.add_paragraph("XGBoost works by building trees one by one. Each new tree tries to correct previous errors. This makes it strong for priority and congestion patterns.")

    doc.add_heading("9. Technology Stack and Responsible Files", level=1)
    add_table(doc, ["Technology", "Used For", "Current Running Files"], [
        ["React", "Builds dashboard UI and tabs.", "frontend/src/main.jsx, App.jsx, components/FlightInputEngine.jsx"],
        ["Vite", "Runs and builds frontend quickly.", "frontend/vite.config.js, package.json"],
        ["Tailwind CSS", "Styling cards, tabs, forms, buttons, responsive layout.", "frontend/src/index.css, tailwind.config.js"],
        ["Zustand", "Stores flights, runways, algorithms, events, weather, predictions.", "frontend/src/store/useTelemetryStore.js"],
        ["Lucide React", "Icons for aircraft, radar, database, clock, alerts.", "FlightInputEngine.jsx"],
        ["Flask", "Backend API for prediction, saving flights, saving schedules, history.", "python_ai/api_server.py"],
        ["Pandas", "Creates ML feature DataFrame.", "python_ai/api_server.py"],
        ["Joblib", "Loads saved ML model .pkl files.", "python_ai/api_server.py, train_rf.py, train_xgboost.py"],
        ["scikit-learn", "Random Forest training and metrics.", "python_ai/pipeline/train_rf.py"],
        ["XGBoost", "Priority and congestion model training.", "python_ai/pipeline/train_xgboost.py"],
        ["C++17", "Fast scheduling engine concepts and classes.", "cpp_scheduler, SchedulerEngine/include"],
        ["CMake", "Builds C++ executable.", "cpp_scheduler/CMakeLists.txt, SchedulerEngine/CMakeLists.txt"],
        ["MySQL Database", "Final project database for persistence.", "python_ai/db_config.py, init_db.py, db_service.py"],
        ["CSV / Parquet datasets", "Training and demo data source.", "Datasets folder"],
    ])

    doc.add_heading("10. Dataset and Database", level=1)
    bullets(doc, [
        "Datasets folder contains aircraft metadata, delay datasets, OpenSky telemetry, cleaned parquet files, merged dataset, and feature datasets.",
        "Dataset is used to train ML models and to simulate batch flight loading.",
        "Main dataset paths include Datasets/aircraft_metadata, Datasets/delays, Datasets/opensky_states, Datasets/cleaned, Datasets/merged_dataset, and Datasets/features.",
        "Database is explained as MySQL for the project. It stores important operational records so the system can show history and archive later.",
    ])
    add_table(doc, ["Database Table", "Purpose"], [
        ["flights", "Stores submitted flight details."],
        ["delay_predictions", "Stores predicted delay, conflict risk, priority score, confidence, and recommended runway."],
        ["schedule_results", "Stores assigned runway, slot time, delay, status, and algorithm."],
        ["runway_states", "Stores runway status, assigned aircraft, timer, wake separation, utilization, and queue length."],
        ["algorithm_switch_log", "Stores old algorithm, new algorithm, and reason for switching."],
        ["event_history", "Stores all important UI/backend events."],
        ["emergencies", "Stores emergency flight records."],
    ])


def teacher_section(doc):
    doc.add_heading("11. How to Explain This Project to Teacher", level=1)
    doc.add_paragraph("Simple explanation script:")
    numbered(doc, [
        "This project is an Airport Runway Scheduler that assigns safe runway slots to landing and takeoff flights.",
        "The user can enter required flight details manually or load dataset/demo flights.",
        "The system calculates priority, delay, risk, and assigns runways using scheduling algorithms.",
        "Default baseline is FCFS, but the system can switch to Priority, Emergency, Greedy, or Hybrid Adaptive AI depending on conditions.",
        "Emergency flights are handled first using Preemptive Emergency logic.",
        "High traffic uses Greedy Min Delay to reduce runway delay.",
        "High storm/risk uses Hybrid Adaptive AI.",
        "Runway Status tab shows live runway occupation, timer, queue, and next flight scheduling messages.",
        "Python ML predicts delay/risk and C++ represents fast scheduling engine concepts.",
        "MySQL stores flights, predictions, schedules, runway states, event logs, and emergency records."
    ])

    doc.add_heading("12. FAQs for Examiner", level=1)
    faqs = [
        ("What is the main aim of this project?", "To schedule airport runway usage safely and efficiently for landing and takeoff flights."),
        ("Which algorithm is default?", "FCFS is the default baseline. The state starts as AUTO, but AUTO behaves like FCFS under normal conditions."),
        ("Why do we need multiple algorithms?", "Different airport situations need different scheduling behavior: normal, priority, emergency, congestion, or risky weather."),
        ("When does Preemptive Emergency run?", "When a flight has medical or technical distress, or emergency flag is detected."),
        ("When does Greedy Min Delay run?", "When traffic is high, flight count is large, or the user clicks Optimize Runways."),
        ("When does Hybrid Adaptive AI run?", "When storm severity or conflict risk is high."),
        ("What is Schedule Time?", "It is expected arrival/takeoff time in simulation seconds."),
        ("What is Runway Use Time?", "It is how long the aircraft occupies the runway."),
        ("What is Safety Gap?", "It is minimum separation time required between aircraft for safety."),
        ("What is Wake Size?", "It is aircraft wake turbulence category: Light, Medium, or Heavy."),
        ("What does Fuel Urgency do?", "It increases priority score. Higher fuel urgency gets more importance."),
        ("What does Storm Severity do?", "It increases weather risk and delay. High storm may select Hybrid Adaptive AI."),
        ("What does Traffic Level do?", "High traffic can trigger Greedy Min Delay scheduling."),
        ("What is Conflict Risk?", "It estimates chance of unsafe or difficult scheduling situation."),
        ("How is priority score calculated?", "priorityScore = emergency + fuelUrgency*9 + passengerWeight + cargoWeight."),
        ("How is runway selected?", "If auto is enabled, the system chooses the runway with lower current load. Otherwise it uses preferred runway if valid."),
        ("What happens after Run Scheduler?", "Flights are ordered, runway slots are assigned, runway board updates, and results are saved/logged."),
        ("How does live runway timer work?", "tickRunways() decreases runway timer every second. When it reaches zero, next queued flight starts or runway becomes free."),
        ("Which ML model is used?", "Random Forest for delay prediction and XGBoost for priority/congestion intelligence."),
        ("What are ML inputs?", "geoaltitude, velocity, wake factor, congestion, runway occupancy, storm, visibility, urgency, and emergency flags."),
        ("What is ML output?", "Predicted delay, risk, priority score, confidence, and recommended runway."),
        ("What if ML model fails?", "Fallback formula in frontend/backend still predicts delay and risk."),
        ("Which database is used?", "MySQL is used in final project explanation for storing flights, predictions, schedules, runway states, event logs, and emergencies."),
        ("Which frontend framework is used?", "React with Vite."),
        ("Which state library is used?", "Zustand."),
        ("Why use C++?", "C++ is fast and suitable for queues, heaps, conflict detection, and scheduling engine logic."),
        ("Which OOP concepts are used?", "Class, object, encapsulation, inheritance, polymorphism, abstraction, operator overloading, strategy pattern, observer pattern."),
        ("Which OS concepts are used?", "Process launching, timers, concurrency/mutex, and inter-process communication through HTTP API."),
        ("Where are flights stored in frontend?", "In flights array inside useTelemetryStore.js."),
        ("Where are schedule logs shown?", "Scheduled Flight Logs card in Simulation Control tab."),
    ]
    add_table(doc, ["Question", "Answer"], faqs)


def main():
    doc = Document()
    setup(doc)
    title(doc)
    dashboard_section(doc)
    simulation_section(doc)
    flow_section(doc)
    algorithm_section(doc)
    concepts_section(doc)
    ml_tech_section(doc)
    teacher_section(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
