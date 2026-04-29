from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
CLASS_JPG = ROOT / "C++_Class_Diagram.jpg"
ARCH_JPG = ROOT / "System_Architecture_Diagram.jpg"
DOCX = ROOT / "AeroSlot_Tech_Stack_Cpp_Project_Report.docx"


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


F_TITLE = font(46, True)
F_H = font(27, True)
F_M = font(22, True)
F_B = font(19)
F_S = font(16)
F_XS = font(14)


def wrap_text(draw, text, fnt, max_width):
    words = str(text).split()
    lines, cur = [], ""
    for word in words:
        trial = word if not cur else cur + " " + word
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines


def rounded_box(draw, xy, fill, outline="#b7d5e8", width=2, radius=22):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, xy, text, fnt, fill="#082f49"):
    x1, y1, x2, y2 = xy
    bbox = draw.textbbox((0, 0), text, font=fnt)
    draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2), text, font=fnt, fill=fill)


def arrow(draw, start, end, color="#1f6f99", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 13
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
        (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill=color)


def class_box(draw, x, y, w, h, title, attrs, methods, fill="#f7fbff"):
    rounded_box(draw, (x, y, x + w, y + h), fill=fill, outline="#8cbdd7", radius=18)
    draw.rounded_rectangle((x, y, x + w, y + 46), radius=18, fill="#0f5f83", outline="#0f5f83")
    center_text(draw, (x, y, x + w, y + 46), title, F_M, fill="white")
    yy = y + 58
    draw.text((x + 16, yy), "Attributes", font=F_S, fill="#0f5f83")
    yy += 24
    for a in attrs:
        for line in wrap_text(draw, "+ " + a, F_XS, w - 30):
            draw.text((x + 16, yy), line, font=F_XS, fill="#123047")
            yy += 18
    yy += 4
    draw.line((x + 12, yy, x + w - 12, yy), fill="#c8ddeb", width=2)
    yy += 10
    draw.text((x + 16, yy), "Methods / Role", font=F_S, fill="#0f5f83")
    yy += 24
    for m in methods:
        for line in wrap_text(draw, "+ " + m, F_XS, w - 30):
            draw.text((x + 16, yy), line, font=F_XS, fill="#123047")
            yy += 18


def generate_class_diagram():
    img = Image.new("RGB", (2600, 1850), "#f3f8fc")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "AeroSlot C++ Class Diagram", font=F_TITLE, fill="#082f49")
    d.text((72, 105), "Core C++ scheduler classes, strategies, safety managers, and data structures used in the runway scheduling engine", font=F_B, fill="#436276")

    class_box(d, 80, 180, 360, 310, "Aircraft / Flight",
              ["flight_id, arrival_time", "runway_occupancy, wake_category", "priority_score, congestion_level", "emergency flags, waiting_time_factor"],
              ["computeETA()", "updateState()", "calculateRunwayOccupancy()", "getEffectivePriority()"],
              "#fff")
    class_box(d, 520, 180, 390, 310, "ISchedulerStrategy",
              ["abstract interface"],
              ["addAircraft()", "extractNext()", "isEmpty()", "virtual destructor"],
              "#eef8ff")
    class_box(d, 80, 600, 360, 300, "Scheduler",
              ["unique_ptr<ISchedulerStrategy>", "SchedulerPolicyManager", "RunwayResourceManager", "ConflictDetector", "ConstraintManager"],
              ["generateStrategyEnvelope()", "addPendingAircraft()", "processNext()", "onEmergency()"],
              "#fff")
    class_box(d, 520, 600, 360, 300, "SchedulerPolicyManager",
              ["system_load", "context"],
              ["Emergency -> Priority", "load > 0.85 -> Greedy", "load > 0.50 -> Priority", "else -> FCFS"],
              "#fff")
    class_box(d, 980, 120, 340, 250, "FCFSStrategy",
              ["queue<Aircraft> q"],
              ["push in arrival order", "pop front aircraft"],
              "#f7fff8")
    class_box(d, 980, 420, 340, 280, "PriorityStrategy",
              ["priority_queue", "PriorityAgingCompare"],
              ["priority + waiting time", "arrival time as fallback"],
              "#f7fff8")
    class_box(d, 980, 750, 340, 280, "GreedyStrategy",
              ["priority_queue candidate_pool", "GreedyDelayCompare"],
              ["penalty = separation*2 - priority", "min-delay best-fit extraction"],
              "#f7fff8")
    class_box(d, 1420, 180, 370, 310, "ConflictDetector",
              ["IntervalTree local_tactical_tree", "runway interval schedules"],
              ["is_safe(candidate)", "commit_allocation()", "clear_path_for()", "overlap prevention"],
              "#fff")
    class_box(d, 1420, 600, 370, 300, "IntervalTree",
              ["IntervalNode root", "start, end, max_end", "left/right nodes"],
              ["addInterval()", "checkOverlap()", "O(log n)-style interval search"],
              "#fff")
    class_box(d, 1880, 180, 360, 300, "ConstraintManager",
              ["unique_ptr<WakeSeparationManager>"],
              ["evaluateOperationalConstraints()", "delegates wake safety checks"],
              "#fff")
    class_box(d, 1880, 570, 360, 310, "WakeSeparationManager",
              ["map<pair<char,char>, int>", "weather_factor"],
              ["getSeparationBuffer()", "updateWeatherFactor()", "check_wake_separation()"],
              "#fff")
    class_box(d, 1880, 960, 360, 290, "RunwayResourceManager",
              ["bool is_locked", "optional<Aircraft> occupant", "mutex mtx"],
              ["check_runway_availability()", "assign()", "release()", "reservation_override()"],
              "#fff")
    class_box(d, 1420, 1020, 370, 300, "EmergencyManager",
              ["vector<IEmergencyObserver*> observers"],
              ["registerObserver()", "compareEmergencies()", "handleSystemInterrupt()"],
              "#fff7f7")
    class_box(d, 520, 1020, 360, 290, "Runway",
              ["runway_id, runway_type", "status, release_time", "is_locked"],
              ["isAvailable()", "lockRunway()", "unlockRunway()"],
              "#fff")
    class_box(d, 80, 1020, 360, 290, "SmartScheduler",
              ["queue fcfs_queue", "priority_queue priority_heap", "ConflictDetector safety_detector"],
              ["add_flight()", "schedule_fcfs()", "schedule_priority()", "attempt_allocation()"],
              "#eef8ff")

    arrow(d, (440, 335), (520, 335))
    arrow(d, (900, 335), (980, 245))
    arrow(d, (900, 355), (980, 560))
    arrow(d, (900, 375), (980, 890))
    arrow(d, (440, 750), (520, 750))
    arrow(d, (880, 740), (980, 560))
    arrow(d, (440, 740), (1420, 330))
    arrow(d, (1790, 335), (1880, 330))
    arrow(d, (1600, 490), (1600, 600))
    arrow(d, (1790, 740), (1880, 725))
    arrow(d, (440, 1165), (520, 1165))
    arrow(d, (1790, 1170), (1880, 1110))
    arrow(d, (80 + 180, 1020), (80 + 180, 900))
    arrow(d, (260, 600), (260, 490))

    d.text((90, 1405), "Key relationships", font=F_H, fill="#082f49")
    notes = [
        "Scheduler owns the active strategy through unique_ptr and changes it according to load or emergency context.",
        "FCFSStrategy, PriorityStrategy, and GreedyStrategy implement ISchedulerStrategy, demonstrating runtime polymorphism.",
        "ConflictDetector uses IntervalTree/interval lists to reject overlapping runway slots before allocation.",
        "ConstraintManager calls WakeSeparationManager to enforce wake turbulence buffers affected by aircraft category and weather.",
        "EmergencyManager notifies Scheduler through IEmergencyObserver so emergency flights can preempt normal sequencing.",
        "SmartScheduler is the compact C++ demonstration layer used for FCFS, priority heap, and allocation examples."
    ]
    yy = 1450
    for note in notes:
        d.text((110, yy), "- " + note, font=F_B, fill="#123047")
        yy += 42

    img.save(CLASS_JPG, quality=95)


def layer(draw, idx, y, title, items, fill="#ffffff", accent="#0f6f8f"):
    x, w, h = 210, 2050, 150
    rounded_box(draw, (x, y, x + w, y + h), fill=fill, outline="#b6d5e5", radius=28)
    draw.ellipse((60, y + 42, 130, y + 112), fill=accent)
    center_text(draw, (60, y + 42, 130, y + 112), str(idx), F_H, fill="white")
    draw.text((230, y + 22), title, font=F_H, fill="#08445d")
    item_w = 310
    start = 245
    for i, item in enumerate(items):
        ix = start + i * item_w
        rounded_box(draw, (ix, y + 70, ix + item_w - 25, y + 132), fill="#f6fbff", outline="#d3e6f0", radius=14)
        center_text(draw, (ix + 8, y + 70, ix + item_w - 33, y + 132), item, F_S, fill="#123047")


def generate_architecture_diagram():
    img = Image.new("RGB", (2400, 2500), "#f4f9fd")
    d = ImageDraw.Draw(img)
    d.text((110, 55), "AeroSlot Scheduler System Architecture", font=F_TITLE, fill="#082f49")
    d.text((112, 115), "Hybrid React + Python AI + C++ scheduling engine with MySQL persistence", font=F_B, fill="#436276")

    layers = [
        ("User / Real Inputs", ["Flight details", "Arrival/Departure", "Priority flags", "Weather", "Congestion", "Dataset load"]),
        ("React Frontend", ["Flight Intake", "Simulation Control", "Decision Tower", "Live Operations", "Archive Analytics"]),
        ("Zustand + Gateway Logic", ["State store", "Prediction fallback", "Algorithm chooser", "Fetch API", "Event logging"]),
        ("Python AI / Flask API", ["RF delay model", "XGBoost priority", "XGBoost congestion", "Feature builder", "Safe fallback"]),
        ("Hybrid C++ Scheduling Engine", ["FCFS queue", "Priority heap", "Emergency preemption", "Greedy min-delay", "Adaptive policy"]),
        ("Safety Simulation Layer", ["Wake separation", "Interval overlap", "Runway locks", "Conflict detector", "Resource manager"]),
        ("Final Schedule Output", ["Landing sequence", "Takeoff sequence", "Runway assignment", "Delay minimized", "Alerts"]),
        ("MySQL Persistence Layer", ["flights", "delay_predictions", "schedule_results", "runway_states", "algorithm_switch_log", "event_history", "emergencies"]),
        ("Archive + Feedback Learning", ["Historical analytics", "KPIs", "Delay statistics", "Conflict outcomes", "Model retraining feedback"]),
    ]
    y = 190
    for idx, (title, items) in enumerate(layers, start=1):
        layer(d, idx, y, title, items, fill="#ffffff" if idx % 2 else "#fbfdff")
        if idx < len(layers):
            arrow(d, (1235, y + 150), (1235, y + 205))
        y += 205

    rounded_box(d, (1680, 735, 2265, 915), "#fff7f7", "#f5b7b1", radius=24)
    d.text((1710, 760), "Safe-mode fallback", font=F_M, fill="#9b1c1c")
    for i, txt in enumerate(["Used when ML confidence is low", "API/model fails", "emergency needs deterministic scheduling"]):
        d.text((1720, 805 + i * 32), "- " + txt, font=F_S, fill="#552020")
    arrow(d, (1680, 820), (1430, 1015), color="#b42318")

    rounded_box(d, (1850, 1850, 2265, 2240), "#f0fbf7", "#7bc8a4", radius=24)
    d.text((1880, 1880), "Continuous feedback loop", font=F_M, fill="#0b6b4a")
    for i, txt in enumerate(["events and schedules saved", "analytics evaluates results", "new data improves features/models"]):
        d.text((1890, 1930 + i * 35), "- " + txt, font=F_S, fill="#174434")
    arrow(d, (1850, 2035), (2260, 1870), color="#0f7a55")
    arrow(d, (2260, 1870), (2260, 325), color="#0f7a55")
    arrow(d, (2260, 325), (2260, 190), color="#0f7a55")

    d.text((115, 2290), "Flow summary: user inputs enter React, features are prepared, Python ML predicts delay/risk, C++ chooses a safe schedule, MySQL stores all results, and analytics feeds future improvement.", font=F_B, fill="#123047")
    img.save(ARCH_JPG, quality=95)


def set_doc_defaults(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style].font.name = "Calibri"
        styles[style].font.bold = True


def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "24")
        elem.set(qn("w:color"), "9CC7DA")
        pgBorders.append(elem)
    sectPr.append(pgBorders)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AeroSlot Scheduler Project Documentation")
    run.bold = True
    run.font.size = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tech Stack, Architecture, C++ Concepts, ML Models, Algorithms, Formulas, and Examiner FAQs").bold = True
    doc.add_paragraph("Project type: Hybrid AI + C++ runway scheduling system")
    doc.add_paragraph("Database used: MySQL, as specified for the project documentation.")
    doc.add_picture(str(ARCH_JPG), width=Inches(6.6))
    doc.add_page_break()


def generate_docx():
    doc = Document()
    set_doc_defaults(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        add_page_border(section)
    add_title(doc)

    doc.add_heading("1. Tech Stack Used and Main Logic", level=1)
    add_table(doc, ["Layer / Technology", "Purpose", "Main Logic in Simple Language"], [
        ["React 19 + Vite", "Frontend dashboard", "Shows tabs for flight input, simulation, decision tower, live operations, and archive analytics."],
        ["Tailwind CSS", "Styling", "Creates responsive cards, buttons, forms, runway boards, and timeline layouts."],
        ["Zustand", "Frontend state management", "Keeps flights, predictions, selected algorithm, runways, events, and KPIs in one shared store."],
        ["Recharts", "Charts", "Displays algorithm comparison and delay history charts."],
        ["Lucide React", "Icons", "Adds visual icons for aircraft, radar, database, alerts, and metrics."],
        ["Flask API", "Python backend service", "Receives predict/save/history requests from the frontend and returns ML decisions."],
        ["scikit-learn Random Forest", "Delay prediction model", "Predicts expected flight delay from engineered aircraft and environment features."],
        ["XGBoost", "Priority/congestion models", "Learns priority/congestion patterns and supports smarter scheduling decisions."],
        ["Joblib", "Model loading/saving", "Stores trained ML models as .pkl files and loads them in the API."],
        ["Pandas", "Feature frame creation", "Builds tabular input features for ML models."],
        ["C++17", "Scheduling engine", "Implements fast deterministic scheduling using queues, heaps, interval checks, strategy pattern, and emergency interrupts."],
        ["CMake", "C++ build system", "Compiles the C++ scheduler project."],
        ["MySQL", "Persistence layer", "Stores flights, predictions, schedule results, runway states, switch logs, events, and emergencies."],
        ["Datasets / Parquet / CSV", "Training and batch data", "Stores cleaned telemetry, delay, aircraft metadata, and merged training records."],
    ])

    doc.add_heading("2. Important Files and Their Features", level=1)
    add_table(doc, ["File / Folder", "Technology", "What It Does"], [
        ["launch_system.py", "Python subprocess", "Starts the Python service, C++ scheduler executable, and React frontend together."],
        ["frontend/src/main.jsx", "React", "Mounts the React app into the browser."],
        ["frontend/src/App.jsx", "React", "Loads the main FlightInputEngine component."],
        ["frontend/src/components/FlightInputEngine.jsx", "React + UI", "Contains all main tabs: Flight Intake, Simulation Control, Decision Tower, Live Operations, and Archive Analytics."],
        ["frontend/src/store/useTelemetryStore.js", "Zustand + JS logic", "Main frontend brain. It stores flights, predicts delay, chooses algorithms, schedules flights, updates runway states, computes KPIs, and logs events."],
        ["frontend/src/services/bffGateway.ts", "TypeScript + Axios", "Gateway idea for routing frontend requests to C++ and Python services through a backend-for-frontend."],
        ["frontend/src/components/HistoryLogs.jsx", "React", "Displays previous events/history information."],
        ["frontend/src/store/sqliteService.js", "JS service placeholder", "Persistence helper naming remains old, but project documentation should describe the actual database as MySQL."],
        ["python_ai/api_server.py", "Flask + ML", "Exposes /api/predict, /api/flights, /api/schedule, /api/history, /api/archive and saves data."],
        ["python_ai/db_config.py", "Database config", "Database connection file. For final project explanation, configure this for MySQL connection details."],
        ["python_ai/init_db.py", "SQL schema", "Creates tables such as flights, schedule_results, delay_predictions, runway_states, event_history, and emergencies."],
        ["python_ai/db_service.py", "Database service", "Contains insert/select functions for flights, predictions, assignments, switches, logs, runway states, and emergencies."],
        ["python_ai/pipeline/train_rf.py", "scikit-learn", "Trains RandomForestRegressor for delay prediction and saves rf_delay_model.pkl."],
        ["python_ai/pipeline/train_xgboost.py", "XGBoost", "Trains XGBoost models for priority and congestion intelligence."],
        ["python_ai/pipeline/data_loader.py", "Python data pipeline", "Loads datasets for model training/evaluation."],
        ["python_ai/pipeline/evaluate_models.py", "ML evaluation", "Evaluates model quality using validation metrics."],
        ["python_ai/model_registry.py", "Model management", "Central location for managing available model artifacts."],
        ["cpp_scheduler/CMakeLists.txt", "CMake", "Builds the C++17 airport_scheduler executable."],
        ["cpp_scheduler/include/models/Aircraft.h", "C++ class", "Defines aircraft fields like id, ETA, emergency flag, wake category, and ML delay score."],
        ["cpp_scheduler/include/models/Runway.h", "C++ class", "Represents runway status, lock state, and release time."],
        ["cpp_scheduler/include/strategies/Scheduler.h", "C++ scheduler", "Defines SmartScheduler with FCFS queue, priority heap, conflict detector, and allocation methods."],
        ["cpp_scheduler/include/strategies/ConflictDetector.h", "C++ safety", "Stores runway time intervals and checks if requested slots overlap."],
        ["SchedulerEngine/include/Scheduler.h", "C++ orchestration", "Main advanced scheduler that selects strategy, checks constraints, checks runway resources, and handles emergencies."],
        ["SchedulerEngine/include/ISchedulerStrategy.h", "C++ interface", "Abstract base class for scheduling strategies."],
        ["SchedulerEngine/include/FCFSStrategy.h", "C++ queue", "Schedules flights in first-come-first-served order."],
        ["SchedulerEngine/include/PriorityStrategy.h", "C++ priority queue", "Schedules higher priority and longer waiting flights first."],
        ["SchedulerEngine/include/GreedyStrategy.h", "C++ heap", "Chooses flights that minimize delay penalty."],
        ["SchedulerEngine/include/ConflictDetector.h", "C++ interval tree", "Checks local runway interval conflicts before committing allocations."],
        ["SchedulerEngine/include/IntervalTree.h", "C++ data structure", "Stores allocated time intervals and detects overlaps efficiently."],
        ["SchedulerEngine/include/WakeSeparationManager.h", "C++ safety matrix", "Calculates wake separation buffer based on leader/follower aircraft categories."],
        ["SchedulerEngine/include/RunwayResourceManager.h", "C++ concurrency", "Locks/unlocks runway resource using mutex and optional occupant."],
        ["SchedulerEngine/include/EmergencyManager.h", "C++ observer pattern", "Notifies scheduler when an emergency aircraft interrupts normal flow."],
    ])

    doc.add_heading("3. System Architecture Explanation", level=1)
    doc.add_picture(str(ARCH_JPG), width=Inches(6.6))
    add_bullets(doc, [
        "The user enters flight, weather, priority, congestion, and dataset information in the React frontend.",
        "The frontend state store prepares the data and can call the Python Flask API for ML prediction.",
        "Python AI loads Random Forest and XGBoost models. If models are unavailable, it uses a safe fallback formula.",
        "The scheduler chooses a C++ scheduling style: FCFS, Priority, Preemptive Emergency, Greedy Min Delay, or Hybrid Adaptive.",
        "Before final allocation, safety checks verify wake separation, runway occupancy, and time interval conflicts.",
        "The final output is runway assignment, slot time, delay, emergency handling, and visual timeline.",
        "MySQL stores operational records so the archive and analytics tabs can show previous decisions and performance."
    ])

    doc.add_heading("4. Frontend Tabs and User Inputs", level=1)
    doc.add_heading("4.1 Flight Intake Tab", level=2)
    add_table(doc, ["Input", "What User Should Enter", "Options / Example", "Purpose"], [
        ["Flight ID", "Unique flight code", "AA772, EM901", "Identifies the aircraft in queues and logs."],
        ["Airline", "Airline/operator name", "American, United, Medevac", "Display and record keeping."],
        ["Arrival / Departure", "Whether the aircraft is landing or taking off", "Arrival, Departure", "Decides whether it joins landing or takeoff queue."],
        ["Aircraft Type", "Aircraft model", "A320, B737, B777, A380, E190", "Affects occupancy and operational category."],
        ["Wake Category", "Wake turbulence class", "Light, Medium, Heavy", "Determines required safety separation."],
        ["Origin / Destination", "Airport codes", "LAX, JFK", "Used for flight details and history."],
        ["ETA / ETD", "Expected arrival/departure time in seconds in the simulation", "300", "Base time from which delay is calculated."],
        ["Taxi Time Estimate", "Expected taxi duration", "8", "Used in movement and runway planning."],
        ["Runway Occupancy Time", "How long aircraft occupies runway", "45 seconds", "Used for interval conflict checks."],
        ["Passenger Priority", "Passenger importance category", "Normal, VIP, Medical, High Density", "Adds priority score."],
        ["Cargo Priority", "Cargo urgency category", "Routine, Perishable, Medical, Military", "Adds priority score."],
        ["Fuel Urgency", "Fuel urgency level", "1 to 10", "High value increases priority and can reduce delay."],
        ["Minimum Separation", "Minimum required buffer from other aircraft", "60, 90, 120", "Safety gap for scheduling."],
        ["Requested Runway", "Preferred runway if auto assignment is off", "AUTO, 27L, 27R, 09L, 09R", "Manual runway preference."],
        ["Runway Condition", "Physical runway condition", "DRY, WET, CONTAMINATED", "Affects safety explanation and operational conditions."],
        ["Emergency flags", "Check if medical, technical, or fuel emergency exists", "True/False", "Triggers emergency scheduling behavior."],
        ["Wind, Visibility, Storm Severity, Congestion", "Current environment", "Wind 5, visibility 12, storm 0-10, LOW/MEDIUM/HIGH", "Used by delay/risk formulas and algorithm selection."],
    ])
    add_bullets(doc, [
        "Predict Delay button sends the flight and environment to ML/fallback logic and returns delay, risk, priority score, confidence, and runway recommendation.",
        "Add Flight button saves the aircraft into the active queue and logs a FLIGHT_ADDED event.",
        "Batch Load Dataset loads a demo traffic scenario into the queue.",
        "Load Demo Traffic adds sample flights including an emergency case so the examiner can see scheduling behavior quickly."
    ])

    doc.add_heading("4.2 Simulation Control Tab", level=2)
    add_bullets(doc, [
        "Algorithm cards allow selecting FCFS, Priority, Preemptive Emergency, Greedy Min Delay, or Hybrid Adaptive AI.",
        "Run Scheduler checks the current traffic and environment, chooses/switches algorithm, schedules flights, assigns runways, updates runway board, and logs events.",
        "Compare Algorithms runs all major algorithms on the same flight list and compares average delay.",
        "Optimize Runways forces Greedy Min Delay scheduling to reduce delay and balance runway usage.",
        "Inject Emergency creates an emergency Medevac flight and switches the scheduler to Preemptive Emergency.",
        "Inject Storm changes weather to high storm, low visibility, wet runway, and high congestion."
    ])

    doc.add_heading("4.3 Decision Tower, Live Operations, Archive Tabs", level=2)
    add_bullets(doc, [
        "Decision Tower shows runway states, digital twin movement, chosen algorithm, delay, next aircraft, conflict score, and confidence.",
        "Live Operations shows active queues, KPIs such as average delay, utilization, safety score, flights handled, conflict reduction, and event logs.",
        "Archive / Analytics shows stored history: flights, delay records, algorithm switch count, emergencies, model accuracy, delay trends, and algorithm performance."
    ])

    doc.add_heading("5. Machine Learning Models", level=1)
    add_table(doc, ["Model", "File", "How It Works", "Why It Is Used"], [
        ["Random Forest Regressor", "python_ai/pipeline/train_rf.py and models/rf_delay_model.pkl", "Builds many decision trees and averages their outputs to predict delay seconds.", "Good for tabular aviation features and stable delay estimation."],
        ["XGBoost Priority Classifier", "python_ai/pipeline/train_xgboost.py and models/xgb_priority_model.pkl", "Uses boosted decision trees to learn which flights should be treated as higher priority.", "Helps rank flights based on urgency and operational patterns."],
        ["XGBoost Congestion Regressor", "python_ai/pipeline/train_xgboost.py and models/xgb_congestion_model.pkl", "Predicts congestion impact using gradient boosting.", "Helps identify high-load situations where FCFS may not be enough."],
    ])
    add_bullets(doc, [
        "The API builds features such as geoaltitude, velocity, wake factor, airspace congestion, and runway occupancy estimate.",
        "If model loading fails, api_server.py uses local_prediction so the system still works safely during a demo.",
        "The frontend also has a local estimatePrediction fallback so the UI remains usable even if the backend is offline."
    ])

    doc.add_heading("6. C++ Concepts Used", level=1)
    doc.add_picture(str(CLASS_JPG), width=Inches(6.6))
    add_table(doc, ["C++ Concept", "File(s)", "How It Is Used", "Purpose"], [
        ["Classes and Structs", "Aircraft.h, Runway.h, Scheduler.h", "Aircraft, Runway, Scheduler, managers, and strategy classes model real airport objects.", "Encapsulates scheduling data and behavior."],
        ["Encapsulation", "Runway.h, RunwayResourceManager.h", "Runway fields are private and changed through methods; runway locking is controlled by manager functions.", "Prevents unsafe direct state changes."],
        ["Inheritance", "ISchedulerStrategy.h, FCFSStrategy.h, PriorityStrategy.h, GreedyStrategy.h", "Strategies inherit from ISchedulerStrategy.", "Allows common scheduling interface."],
        ["Polymorphism", "Scheduler.h", "Scheduler stores unique_ptr<ISchedulerStrategy> and calls extractNext without knowing exact strategy.", "Allows dynamic algorithm switching."],
        ["Abstract Class / Interface", "ISchedulerStrategy.h, IEmergencyObserver.h", "Defines required methods for strategies and emergency listeners.", "Supports clean architecture and extendability."],
        ["STL queue", "FCFSStrategy.h, SmartScheduler", "Stores aircraft in arrival order.", "Implements FCFS scheduling."],
        ["STL priority_queue", "PriorityStrategy.h, GreedyStrategy.h", "Ranks aircraft by priority or min-delay penalty.", "Implements fast priority and greedy scheduling."],
        ["Custom Comparators", "PriorityStrategy.h, GreedyStrategy.h, cpp_scheduler Scheduler.h", "Comparator structs control heap ordering.", "Defines what aircraft is extracted first."],
        ["map", "WakeSeparationManager.h, ConflictDetector.h", "Stores wake separation matrix and runway interval schedules.", "Fast structured lookup."],
        ["unique_ptr", "Scheduler.h, ConstraintManager.h", "Owns active strategy and wake manager.", "Automatic memory management."],
        ["optional", "RunwayResourceManager.h", "Represents whether a runway has a current occupant.", "Avoids invalid placeholder objects."],
        ["mutex / lock_guard", "RunwayResourceManager.h", "Locks runway state updates.", "Prevents race conditions in concurrent environments."],
        ["Operator Overloading", "cpp_scheduler/src/models/Aircraft.cpp", "Aircraft::operator< defines ordering behavior.", "Allows aircraft comparison in heap-like structures."],
        ["Interval Tree", "IntervalTree.h", "Stores intervals with max_end and checks overlap.", "Efficient conflict detection for runway slots."],
        ["Observer Pattern", "EmergencyManager.h, Scheduler.h", "EmergencyManager notifies Scheduler via IEmergencyObserver.", "Supports emergency interruption."],
        ["Factory Pattern", "AircraftFactory.h", "Central aircraft creation concept.", "Creates aircraft objects consistently."],
    ])

    doc.add_heading("7. Main Mathematical Formulas", level=1)
    add_table(doc, ["Formula", "Meaning of Parameters", "Used For"], [
        ["priorityScore = emergency + fuelUrgency*9 + passengerWeight + cargoWeight", "emergency=100 if distress, fuelUrgency=1-10, passenger/cargo weights depend on selected option", "Ranks aircraft for priority scheduling."],
        ["delay = max(5, 45 + congestion + weather + wake - urgencyCredit)", "congestion depends LOW/MEDIUM/HIGH, weather uses storm and visibility, wake depends aircraft wake category, urgencyCredit uses fuel/emergency", "Frontend fallback delay prediction."],
        ["risk = clamp(0.08, 0.96, delay/150 + congestionRisk + emergencyRisk)", "delay is predicted delay, congestionRisk added under HIGH congestion, emergencyRisk added for emergency", "Conflict risk display and hybrid algorithm switching."],
        ["slotTime = max(ETA, runwayLoad[runway] + separation)", "ETA is requested time, runwayLoad is when runway becomes free, separation is safety buffer", "Finds safe runway slot."],
        ["delayAfterScheduling = max(0, slotTime - ETA) + risk*8", "slotTime is assigned schedule time, ETA is requested time, risk is conflict probability", "Shows final assigned delay."],
        ["runwayLoad = slotTime + occupancy", "occupancy is runway occupation duration", "Updates runway availability after assignment."],
        ["effectivePriority = priority_score + waiting_time_factor", "priority_score is base urgency, waiting_time_factor grows while waiting", "Avoids starvation in priority queue."],
        ["greedyPenalty = separation_requirement*2 - priority_score", "higher separation increases cost, higher priority reduces cost", "Greedy min-delay ordering."],
        ["interval overlap if max(start1,start2) < min(end1,end2)", "two runway time intervals", "Conflict detector safety check."],
        ["wakeBuffer = baseWakeMatrix[leader, follower] * weather_factor", "leader/follower wake categories and weather multiplier", "Wake separation safety."],
        ["utilization = average runway utilization percentage", "runway board utilization values across runways", "KPI display."],
    ])

    doc.add_heading("8. How Algorithms Are Scheduled", level=1)
    add_table(doc, ["Situation", "Algorithm Scheduled", "How It Works"], [
        ["Any medical, technical, or fuel emergency exists", "PREEMPTIVE_EMERGENCY", "Emergency aircraft is moved ahead. Scheduler reserves/clears runway path and processes it first."],
        ["High congestion or many flights in queue", "GREEDY_MIN_DELAY", "Flights are ordered to minimize delay and runway load. The least costly aircraft is selected first."],
        ["High conflict risk or storm severity", "HYBRID_ADAPTIVE", "Uses ML risk/priority style information and balances priority against ETA."],
        ["Medium load", "PRIORITY", "Priority queue schedules high score and long-waiting aircraft before normal aircraft."],
        ["Stable normal traffic", "FCFS", "Aircraft are handled in the same order they arrived/requested service."],
        ["Manual optimization button clicked", "GREEDY_MIN_DELAY", "Forces heap-based runway optimization regardless of current selected algorithm."],
        ["Emergency button clicked", "PREEMPTIVE_EMERGENCY", "Creates an emergency flight and immediately switches scheduling mode."],
        ["Compare button clicked", "All algorithms", "Runs FCFS, Priority, Emergency, Greedy, and Hybrid on current flights and compares average delay."],
    ])

    doc.add_heading("9. Examiner FAQs", level=1)
    add_table(doc, ["Question", "Simple Answer"], [
        ["What problem does this project solve?", "It allocates safe runway slots for arrivals and departures while reducing delay and handling emergencies."],
        ["Why use C++?", "C++ is fast and deterministic, so it is suitable for scheduling, priority queues, interval checks, and runway safety logic."],
        ["Why use Python?", "Python is used for ML training, model loading, prediction API, and database service code."],
        ["Why use React?", "React provides an interactive dashboard for input, simulation, runway state, analytics, and logs."],
        ["Which database is used?", "MySQL is used for storing flights, predictions, schedule results, runway states, algorithm switch logs, events, and emergencies."],
        ["What are the two main ML models?", "Random Forest predicts delay, and XGBoost supports priority/congestion intelligence."],
        ["What happens if ML fails?", "The API and frontend have fallback formulas, so the scheduler can still run safely."],
        ["How is safety maintained?", "Wake separation buffers, runway occupancy intervals, conflict detection, and emergency override logic protect schedule safety."],
        ["How does emergency handling work?", "Emergency flags trigger preemptive scheduling. Emergency aircraft gets priority and the runway path is cleared/overridden."],
        ["What is FCFS?", "First Come First Served. Aircraft are scheduled according to arrival/request order."],
        ["What is Priority Scheduling?", "Aircraft are ranked using emergency, fuel urgency, passenger priority, cargo priority, and waiting time."],
        ["What is Greedy Min Delay?", "It chooses the aircraft/runway option that gives the lowest immediate delay penalty."],
        ["What is Hybrid Adaptive AI?", "It combines ML risk/priority signals with scheduling logic to adapt under risky or congested conditions."],
        ["What is an interval tree used for?", "It checks whether a new runway time interval overlaps existing allocated intervals."],
        ["What are wake categories?", "Light, Medium, and Heavy aircraft create different turbulence, so following aircraft need different separation times."],
        ["What is the final output?", "Landing/takeoff sequence, assigned runway, slot time, delay, conflict score, events, and stored history."],
    ])

    doc.add_heading("10. Short Conclusion", level=1)
    doc.add_paragraph(
        "AeroSlot Scheduler is a hybrid project where React handles user interaction, Python AI predicts delay and operational risk, C++ performs fast and safe scheduling, and MySQL stores operational records. "
        "The core idea is to select the correct scheduling algorithm based on real conditions: normal traffic uses FCFS, priority traffic uses priority queues, high congestion uses greedy delay minimization, high risk uses hybrid AI logic, and emergencies use preemptive scheduling."
    )
    doc.save(DOCX)


if __name__ == "__main__":
    generate_class_diagram()
    generate_architecture_diagram()
    generate_docx()
    print(CLASS_JPG)
    print(ARCH_JPG)
    print(DOCX)
